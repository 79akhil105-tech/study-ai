import docx
from PIL import Image
import io
from pathlib import Path
from src.parsers.base import BaseParser, ParsedContent
from src.utils.helpers import logger, clean_text

class DocxParser(BaseParser):
    def parse(self, file_path: str) -> ParsedContent:
        text_parts = []
        images = []
        metadata = {}
        try:
            doc = docx.Document(file_path)
            metadata['paragraph_count'] = len(doc.paragraphs)
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            for table in doc.tables:
                rows = [" | ".join(c.text for c in row.cells) for row in table.rows]
                text_parts.append("\n[TABLE]\n" + "\n".join(rows) + "\n[/TABLE]")
            for rel in doc.part.rels.values():
                if "image" in rel.target_ref:
                    try:
                        pil_img = Image.open(io.BytesIO(rel.target_part.blob))
                        if pil_img.mode != 'RGB':
                            pil_img = pil_img.convert('RGB')
                        if pil_img.width > 50 and pil_img.height > 50:
                            images.append(pil_img)
                    except Exception:
                        pass
        except Exception as e:
            logger.error(f"DOCX Parse Error [{file_path}]: {e}")
            raise
        return ParsedContent(text=clean_text("\n".join(text_parts)), images=images, metadata=metadata, source_name=Path(file_path).name)
